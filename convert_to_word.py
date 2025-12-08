"""
Convert REPORT.md to REPORT.docx
Simple markdown to Word converter for the traffic prediction report
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document."""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line, style='Intense Quote')
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.add_run('_' * 80)
        
        # Handle tables
        elif line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            
            # Check if next line is still part of table
            if i + 1 < len(lines) and not lines[i + 1].startswith('|'):
                # Create table
                if len(table_data) > 1:  # Need at least header and one row
                    # Filter out separator rows
                    table_rows = [row for row in table_data if not all('-' in cell for cell in row)]
                    
                    if table_rows:
                        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(table_rows):
                            for col_idx, cell_text in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                
                                # Bold header row
                                if row_idx == 0:
                                    cell.paragraphs[0].runs[0].bold = True
                
                in_table = False
                table_data = []
        
        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            # Skip table separator lines
            if not (line.strip().startswith('|') and all(c in '|-: ' for c in line.strip())):
                text = line.strip()
                
                # Remove markdown formatting but preserve emphasis
                # Bold
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                # Italic
                text = re.sub(r'\*(.*?)\*', r'\1', text)
                # Code
                text = re.sub(r'`(.*?)`', r'\1', text)
                
                if text:
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Empty lines (add spacing)
        else:
            if i > 0 and lines[i-1].strip():  # Only add spacing after content
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"[OK] Successfully converted {md_file} to {docx_file}")


if __name__ == "__main__":
    try:
        convert_markdown_to_docx('REPORT.md', 'REPORT.docx')
        print("\n[SUCCESS] Word document created successfully!")
        print("  File: REPORT.docx")
    except Exception as e:
        print(f"[ERROR] {e}")
        import traceback
        traceback.print_exc()

